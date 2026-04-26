import pdfplumber
import docx
import os


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a PDF, DOCX, or TXT file.
    Returns the extracted text as a string.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_pdf(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            # Also extract tables as text
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = " | ".join(cell or "" for cell in row)
                    text_parts.append(row_text)
    return "\n".join(text_parts)


def _extract_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Extract text from tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)